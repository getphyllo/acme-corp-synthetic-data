"""DOCX toolkit for the Acme report pack (documents 61-80).

Written documents rather than decks: plans, briefs, memos, policy docs, minutes.
Same grounding contract and same palette as lib.Doc / pptx_lib.Deck.

    m = Memo("61-....docx", kicker="ANNUAL OPERATING PLAN",
             title="Acme Corp FY2027 Annual Operating Plan — Narrative & Assumptions",
             subtitle="The revenue build, the cost lines, and the assumptions each one rests on",
             byline="Finance / FP&A · reviewed with the CFO",
             meta=["Period: FY2027 (Jan–Dec 2027)", "Version: v1 · July 2026"],
             short="FY27 AOP narrative")
    m.at_a_glance([("$836M", "FY27 net revenue target"), ...])
    m.h1("THE NUMBER", "1 · What we are committing to")
    m.body("Body text with **inline bold**.")
    m.bullets([...]); m.table(headers, rows); m.image(png, "caption")
    m.callout("Risk", "..."); m.recommendations([(action, owner, when)])
    m.build()
"""
from __future__ import annotations
import os
import docx
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lib import palette, OUT_DIR

FONT = "Calibri"


def _rgb(name):
    return RGBColor.from_string(palette[name].lstrip("#"))


NAVY, INK, TEAL, GOLD, RUST = (_rgb(k) for k in ("navy", "ink", "teal", "gold", "rust"))
SLATE, WHITE, GREEN, AMBER = _rgb("slate"), _rgb("white"), _rgb("green"), _rgb("amber")
HEX = {k: palette[k].lstrip("#") for k in palette}
FOOTER = ("Internal · Confidential — Acme Corp synthetic demo data. "
          "Fictional; not for real decision-making.")


def _shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


# w:tblPr children must appear in schema order or Word ignores them.
_TBLPR_ORDER = ["w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
                "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
                "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
                "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
                "w:tblDescription"]


def _tblpr(tbl, tag):
    """Replace (or create) a w:tblPr child, inserted at its schema position."""
    pr = tbl._tbl.tblPr
    for old in pr.findall(qn(tag)):
        pr.remove(old)
    el = OxmlElement(tag)
    rank = _TBLPR_ORDER.index(tag)
    for child in pr:
        name = child.tag.split("}")[-1]
        key = "w:" + name
        if key in _TBLPR_ORDER and _TBLPR_ORDER.index(key) > rank:
            child.addprevious(el)
            return el
    pr.append(el)
    return el


def _borders(tbl, hexcolor="C9D2DB", sz=4):
    el = _tblpr(tbl, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), hexcolor)
        el.append(e)


def _fixed_width(tbl, widths, avail=Inches(6.6)):
    """Force Word to honour our column widths instead of auto-fitting to text."""
    tbl.autofit = False
    _tblpr(tbl, "w:tblLayout").set(qn("w:type"), "fixed")
    w = _tblpr(tbl, "w:tblW")
    w.set(qn("w:w"), str(int(avail / 635)))          # EMU -> twips
    w.set(qn("w:type"), "dxa")
    n = len(tbl.columns)
    fr = widths or [1.0 / n] * n
    tot = sum(fr)
    for i, f in enumerate(fr):
        cw = Emu(int(avail * f / tot))
        tbl.columns[i].width = cw
        for row in tbl.rows:
            row.cells[i].width = cw


def _cell_margins(tbl, twips=72):
    el = _tblpr(tbl, "w:tblCellMar")
    for edge, v in (("top", 40), ("left", twips), ("bottom", 40), ("right", twips)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        el.append(e)


class Memo:
    def __init__(self, filename, title, kicker, subtitle, byline, short,
                 meta=None, doc_type="Internal document"):
        self.filename = filename
        self.path = os.path.join(OUT_DIR, filename)
        self.short = short
        self.d = docx.Document()
        cp = self.d.core_properties
        cp.title, cp.author, cp.comments = title, "Acme Corp (synthetic)", doc_type
        st = self.d.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.14
        sec = self.d.sections[0]
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.95)
        sec.top_margin, sec.bottom_margin = Inches(0.9), Inches(0.9)
        self._footer(sec)
        self._masthead(title, kicker, subtitle, byline, meta or [], doc_type)

    # ------------------------------------------------------------- primitives --
    def _p(self, text="", size=10.5, color=None, bold=False, italic=False,
           space_after=6, space_before=0, align=None, leading=1.14, caps=False,
           keep_with_next=False):
        p = self.d.add_paragraph()
        pf = p.paragraph_format
        pf.space_after, pf.space_before = Pt(space_after), Pt(space_before)
        pf.line_spacing = leading
        pf.keep_with_next = keep_with_next
        if align is not None:
            p.alignment = align
        self._runs(p, text.upper() if caps else text, size, color or INK, bold, italic)
        return p

    def _runs(self, p, text, size, color, bold, italic=False):
        for i, chunk in enumerate(str(text).split("**")):
            if not chunk:
                continue
            r = p.add_run(chunk)
            r.font.name, r.font.size, r.font.color.rgb = FONT, Pt(size), color
            r.font.bold = bold or (i % 2 == 1)
            r.font.italic = italic
        return p

    def _footer(self, sec):
        p = sec.footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(FOOTER + "     ")
        r.font.name, r.font.size, r.font.color.rgb = FONT, Pt(7.5), SLATE
        r2 = p.add_run(self.short)
        r2.font.name, r2.font.size, r2.font.color.rgb = FONT, Pt(7.5), SLATE
        r2.font.bold = True

    def _masthead(self, title, kicker, subtitle, byline, meta, doc_type):
        self._p(kicker, size=9.5, color=TEAL, bold=True, caps=True, space_after=3)
        self._p(title, size=23, color=NAVY, bold=True, space_after=6, leading=1.08)
        self._p(subtitle, size=13, color=SLATE, space_after=8, leading=1.2)
        self._rule()
        self._p(byline, size=10, color=SLATE, bold=True, space_after=2)
        for m in meta:
            self._p(m, size=9.5, color=SLATE, space_after=1)
        self._p(doc_type + " · Internal · Confidential — synthetic demo data",
                size=9.5, color=SLATE, space_after=10)

    def _rule(self, color="C9D2DB", size=6):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        bd = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
        bd.append(b)
        pPr.append(bd)

    # ----------------------------------------------------------------- blocks --
    def at_a_glance(self, pairs, label="At a glance"):
        """pairs: list of (value, label). Renders a shaded n-column strip."""
        self.h2(label)
        t = self.d.add_table(rows=1, cols=len(pairs))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _borders(t)
        _cell_margins(t, 96)
        _fixed_width(t, None)
        for i, (v, l) in enumerate(pairs):
            c = t.cell(0, i)
            _shade(c, HEX["mist2"])
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            self._runs(c.paragraphs[0], str(v), 17, NAVY, True)
            p2 = c.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            self._runs(p2, l.upper(), 8, SLATE, True)
        self._p("", size=4, space_after=4)
        return t

    def h1(self, kicker, headline=None):
        if headline is None:
            kicker, headline = None, kicker
        if kicker:
            self._p(kicker, size=9, color=TEAL, bold=True, caps=True,
                    space_before=12, space_after=2, keep_with_next=True)
        self._p(headline, size=15.5, color=NAVY, bold=True,
                space_before=0 if kicker else 12, space_after=5, keep_with_next=True)
        return self

    def h2(self, text):
        self._p(text, size=11.5, color=NAVY, bold=True, space_before=9,
                space_after=3, keep_with_next=True)
        return self

    def body(self, text, size=10.5):
        self._p(text, size=size)
        return self

    def lede(self, text):
        self._p(text, size=12, color=SLATE, italic=True, space_after=8, leading=1.22)
        return self

    def bullets(self, items, numbered=False):
        for i, it in enumerate(items):
            p = self.d.add_paragraph(style="List Number" if numbered else "List Bullet")
            pf = p.paragraph_format
            pf.space_after, pf.line_spacing = Pt(4), 1.14
            pf.left_indent, pf.first_line_indent = Inches(0.28), Inches(-0.16)
            self._runs(p, it, 10.5, INK, False)
        self._p("", size=2, space_after=2)
        return self

    def table(self, headers, rows, widths=None, total_row=False, note=None,
              status_col=None, size=9.5, align_right_from=1):
        t = self.d.add_table(rows=len(rows) + 1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _borders(t)
        _cell_margins(t)
        _fixed_width(t, widths)
        for j, h in enumerate(headers):
            c = t.cell(0, j)
            _shade(c, HEX["navy"])
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            if j >= align_right_from:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._runs(p, str(h), size - 0.5, WHITE, True)
        for i, row in enumerate(rows, start=1):
            last = total_row and i == len(rows)
            for j, v in enumerate(row):
                c = t.cell(i, j)
                _shade(c, HEX["mist"] if last else (HEX["white"] if i % 2 else HEX["mist2"]))
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                if j >= align_right_from:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                col = INK
                if status_col is not None and j == status_col:
                    col = {"On track": GREEN, "Watch": AMBER, "Action": RUST,
                           "Reinvest": GREEN, "Reallocate": RUST, "Pass": GREEN,
                           "Fail": RUST, "Go": GREEN, "No-go": RUST}.get(str(v), INK)
                self._runs(p, "" if v is None else str(v), size, col,
                           bool(last or (status_col is not None and j == status_col)))
        self._p("", size=3, space_after=3)
        if note:
            self.source(note)
        return t

    def callout(self, title, text, kind="info"):
        col = {"info": "teal", "risk": "rust", "action": "gold", "win": "green"}.get(kind, "teal")
        t = self.d.add_table(rows=1, cols=1)
        _cell_margins(t, 110)
        _fixed_width(t, None)
        c = t.cell(0, 0)
        _shade(c, HEX["mist2"])
        el = _tblpr(t, "w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            if edge == "left":
                e.set(qn("w:val"), "single")
                e.set(qn("w:sz"), "24")
                e.set(qn("w:color"), HEX[col])
            else:
                e.set(qn("w:val"), "none")
            el.append(e)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        self._runs(p, title, 10.5, NAVY, True)
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.14
        self._runs(p2, text, 10, INK, False)
        self._p("", size=4, space_after=4)
        return t

    def image(self, png, caption=None, width=6.6):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(png, width=Inches(width))
        if caption:
            self._p(caption, size=8.5, color=SLATE, italic=True, space_after=8)
        return self

    def source(self, text):
        text = str(text)
        if text.lower().startswith("source:"):        # callers pass it either way
            text = text.split(":", 1)[1].strip()
        self._p("Source: " + text, size=8, color=SLATE, italic=True, space_after=8)
        return self

    def recommendations(self, items, headline="Recommendations & next steps"):
        self.h2(headline)
        rows = [[str(i + 1), a, o, w] for i, (a, o, w) in enumerate(items)]
        return self.table(["#", "Action", "Owner", "When"], rows,
                          widths=[0.05, 0.55, 0.24, 0.16], align_right_from=3)

    def decisions(self, items, headline="Decision log"):
        """items: (decision, owner, date, status)"""
        self.h2(headline)
        return self.table(["Decision", "Owner", "Date", "Status"], items,
                          widths=[0.46, 0.22, 0.14, 0.18], status_col=3, align_right_from=2)

    def risks(self, items, headline="Risks & mitigations"):
        """items: (risk, impact, mitigation, owner)"""
        self.h2(headline)
        return self.table(["Risk", "Impact", "Mitigation", "Owner"], items,
                          widths=[0.24, 0.2, 0.38, 0.18], align_right_from=9)

    def signoff(self, rows, headline="Approvals"):
        """rows: (name/role, what they are approving)"""
        self.h2(headline)
        return self.table(["Approver", "Approving", "Date"],
                          [[a, b, "________"] for a, b in rows],
                          widths=[0.3, 0.52, 0.18], align_right_from=2)

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return self

    def build(self):
        self.d.save(self.path)
        return self.path
